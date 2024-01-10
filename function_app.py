import azure.functions as func
import json
import logging
import base64
import io
from docx import Document
import traceback
from docx.oxml.ns import qn
import pandas as pd
import tempfile
import logging
import os
app = func.FunctionApp()

@app.route(route="extract", auth_level=func.AuthLevel.ANONYMOUS)
def extract(req: func.HttpRequest) -> func.HttpResponse:
    try:
        file_bytes = req.get_body()
        if not file_bytes:
            return func.HttpResponse("Request body is empty", status_code=400)

        doc = Document(io.BytesIO(file_bytes))
        extracted_data = extract_data_from_controls(doc)
        logging.info(f"Extracted data: {extracted_data}")

        return func.HttpResponse("Document processed successfully", status_code=200)

    except Exception as e:
        logging.error("Exception type: " + str(type(e)))
        logging.error("Error processing request: " + str(e))
        logging.error("Traceback: " + traceback.format_exc())
        return func.HttpResponse("Error processing request: " + str(e), status_code=500)

@app.route(route="track_changes", methods=['POST'])
def track_changes(req: func.HttpRequest) -> func.HttpResponse:
    try:
        file_bytes = req.get_body()
        if not file_bytes:
            return func.HttpResponse("Request body is empty", status_code=400)

        doc = Document(io.BytesIO(file_bytes))
        extracted_data = extract_data_with_track_changes(doc)
        logging.info(f"Extracted data: {extracted_data}")

        return func.HttpResponse("Document processed successfully", status_code=200)

    except Exception as e:
        logging.error("Exception type: " + str(type(e)))
        logging.error("Error processing request: " + str(e))
        logging.error("Traceback: " + traceback.format_exc())
        return func.HttpResponse("Error processing request: " + str(e), status_code=500)


def extract_data_from_controls(doc):
    data = {}
    for part in doc.element.body.iter():
        if part.tag.endswith('sdt'):
            sdtPr = part.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdtPr')
            tagElem = sdtPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tag') if sdtPr is not None else None
            tag = tagElem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') if tagElem is not None else None
            
            texts = part.xpath('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            if texts:
                content = ''.join(t.text for t in texts)
                if tag:
                    if tag not in data:
                        data[tag] = []
                    data[tag].append(content)

    logging.info(f"Data extracted from document controls: {data}")
    return data
def extract_track_changes(p_element):
    """
    Extracts text including proposed insertions, deletions, and modifications in track changes.
    """
    text = ''
    for child in p_element.iterchildren():
        if child.tag == qn('w:t'):  # Regular text
            text += child.text
        elif child.tag == qn('w:ins'):  # Inserted text
            inserted_text = ''.join(node.text for node in child.iterdescendants(tag=qn('w:t')))
            text += f" [Inserted: {inserted_text}]"
        elif child.tag == qn('w:del'):  # Deleted text
            deleted_text = ''.join(node.text for node in child.iterdescendants(tag=qn('w:t')))
            text += f" [Deleted: {deleted_text}]"
    return text
def extract_data_with_track_changes(doc):
    data = []
    for p in doc.paragraphs:
        text = extract_track_changes(p._element)
        data.append(text)
    return [element for element in data if element != ""]
